import io

import pytest
from docx import Document

from app.services.document_pipeline import (
    build_clean_chunks,
    chunk_text_sentence_aware,
    clean_raw_text,
    detect_language_hint,
    extract_text_from_upload,
    normalize_text_value,
)


class TestTextCleaning:
    def test_clean_raw_text_keeps_vietnamese(self):
        text = "Học sinh    cần ôn tập bài hôm nay"
        cleaned = clean_raw_text(text)
        assert "Học sinh" in cleaned
        assert "  " not in cleaned

    def test_normalize_text_value_ascii(self):
        normalized = normalize_text_value("Học sinh cần ôn tập")
        assert normalized == "hoc sinh can on tap"

    def test_detect_language_hint(self):
        assert detect_language_hint("Đây là tiếng Việt") == "vi"
        assert detect_language_hint("This is English text") == "en"


class TestChunking:
    def test_chunk_short_text(self):
        text = "Đây là đoạn ngắn."
        chunks = chunk_text_sentence_aware(text, chunk_size=100, chunk_overlap=20)
        assert chunks == [text]

    def test_chunk_long_text(self):
        text = "Câu hỏi trắc nghiệm giúp học nhanh. " * 80
        chunks = chunk_text_sentence_aware(text, chunk_size=250, chunk_overlap=50)
        assert len(chunks) > 1

    def test_build_clean_chunks_deduplicates(self):
        text = "Học sinh cần ôn tập. Hoc sinh can on tap. Học sinh cần ôn tập."
        rows = build_clean_chunks(
            raw_text=text,
            chunk_size=120,
            chunk_overlap=30,
            min_chunk_length=5,
        )
        assert len(rows) >= 1
        normalized_values = [row["normalized_text"] for row in rows]
        assert len(normalized_values) == len(set(normalized_values))

    def test_chunk_long_sentence_no_punctuation(self):
        text = "A" * 600
        chunks = chunk_text_sentence_aware(text, chunk_size=120, chunk_overlap=20)
        assert len(chunks) > 1
        assert all(len(chunk) <= 120 for chunk in chunks)


class TestExtractors:
    def test_extract_docx_with_table(self):
        doc = Document()
        doc.add_paragraph("Doan 1.")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "B1"
        table.cell(1, 0).text = "A2"
        table.cell(1, 1).text = "B2"

        buffer = io.BytesIO()
        doc.save(buffer)

        text = extract_text_from_upload(
            file_name="sample.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            content=buffer.getvalue(),
        )

        assert "Doan 1." in text
        assert "A1 | B1" in text

    def test_extract_txt_utf16(self):
        content = "Xin chao".encode("utf-16")
        text = extract_text_from_upload(
            file_name="sample.txt",
            content_type="text/plain",
            content=content,
        )
        assert "Xin chao" in text


@pytest.mark.parametrize(
    "raw,expected",
    [
        ("", ""),
        ("   ", ""),
        ("Q&amp;A", "Q&A"),
    ],
)
def test_clean_raw_text_parametrized(raw, expected):
    assert clean_raw_text(raw) == expected
